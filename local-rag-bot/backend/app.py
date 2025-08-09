from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import uuid
from datetime import datetime
import numpy as np
from sentence_transformers import SentenceTransformer
from pymongo import MongoClient
import PyPDF2
import docx
import requests
import faiss
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx'}
OLLAMA_BASE_URL = "http://localhost:11434"
ALLOWED_MODELS = ['gemma3:1b', 'mistral:latest', 'llama3.2:1b']

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# MongoDB setup
try:
    client = MongoClient('mongodb://localhost:27017/', serverSelectionTimeoutMS=5000)
    client.admin.command('ping')
    db = client['starrag_bot']
    documents_collection = db['documents']
    conversations_collection = db['conversations']
    print("✅ MongoDB connected successfully")
except Exception as e:
    print(f"⚠️ MongoDB connection failed: {e}")
    documents_collection = None
    conversations_collection = None
    client = None

embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
faiss_index = None
document_chunks = []

class RAGPipeline:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
        self.embedding_model = embedding_model
        
    def allowed_file(self, filename):
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

    def extract_text_from_file(self, file_path, filename):
        """Extract text from different file types"""
        text = ""
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        try:
            if file_extension == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    text = file.read()
            elif file_extension == 'pdf':
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            elif file_extension == 'docx':
                doc = docx.Document(file_path)
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting text from {filename}: {str(e)}")
            return None
            
        return text

    def process_document(self, file_path, filename):
        """Process a document and add it to the vector store"""
        global faiss_index, document_chunks
        
        # Extract text
        text = self.extract_text_from_file(file_path, filename)
        if not text:
            return False
            
        # Create document and split into chunks
        doc = Document(page_content=text, metadata={"source": filename})
        chunks = self.text_splitter.split_documents([doc])
        
        # Generate embeddings
        chunk_texts = [chunk.page_content for chunk in chunks]
        embeddings = self.embedding_model.encode(chunk_texts)
        
        # Store in MongoDB
        doc_id = str(uuid.uuid4())
        document_data = {
            "_id": doc_id,
            "filename": filename,
            "content": text,
            "chunks": [
                {
                    "text": chunk.page_content,
                    "metadata": chunk.metadata,
                    "embedding": embedding.tolist()
                }
                for chunk, embedding in zip(chunks, embeddings)
            ],
            "created_at": datetime.utcnow()
        }
        
        if documents_collection is not None:
            try:
                documents_collection.insert_one(document_data)
            except Exception as e:
                print(f"Warning: Could not store document in MongoDB: {e}")
        else:
            print("Warning: MongoDB not available - document not persisted")
        
        # Add to FAISS index
        if faiss_index is None:
            dimension = embeddings.shape[1]
            faiss_index = faiss.IndexFlatL2(dimension)
            document_chunks = []
        
        faiss_index.add(embeddings.astype('float32'))
        
        for i, chunk in enumerate(chunks):
            document_chunks.append({
                "text": chunk.page_content,
                "metadata": chunk.metadata,
                "doc_id": doc_id
            })
        
        return True

    def similarity_search(self, query, k=5):
        """Search for similar chunks using FAISS"""
        global faiss_index, document_chunks
        
        if faiss_index is None or len(document_chunks) == 0:
            return []
        
        # Generate query embedding
        query_embedding = self.embedding_model.encode([query])
        
        # Search in FAISS
        distances, indices = faiss_index.search(query_embedding.astype('float32'), k)
        
        # Return relevant chunks
        results = []
        for i, idx in enumerate(indices[0]):
            if idx < len(document_chunks):
                chunk = document_chunks[idx]
                results.append({
                    "text": chunk["text"],
                    "metadata": chunk["metadata"],
                    "score": float(distances[0][i])
                })
        
        return results

    def generate_response(self, query, context_chunks, model="gemma3:1b"):
        """Generate response using Ollama"""
        context = "\n\n".join([chunk["text"] for chunk in context_chunks])
        
        prompt = f"""Based on the following context, answer the question. If you can't find the answer, say so.

Context:
{context}

Question: {query}

Answer:"""

        try:
            # Check Ollama health
            health_check = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=30)
            if health_check.status_code != 200:
                return "Error: Ollama is not responding. Please run 'ollama serve'."
            
            # Verify model is available
            models = health_check.json().get('models', [])
            model_names = [m.get('name', '') for m in models]
            
            if model not in model_names:
                # Try to find an available allowed model
                available_models = [m for m in ALLOWED_MODELS if m in model_names]
                if available_models:
                    model = available_models[0]
                    print(f"Using available model: {model}")
                else:
                    available = ', '.join(model_names) if model_names else 'none'
                    return f"Error: None of the allowed models are available. Available models: {available}"
            
            # Generate response
            response = requests.post(
                f"{OLLAMA_BASE_URL}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "temperature": 0.7,
                        "top_p": 0.9,
                        "num_predict": 2048
                    }
                },
                timeout=300
            )
            
            if response.status_code == 200:
                return response.json().get("response", "Sorry, I couldn't generate a response.")
            else:
                error_msg = f"Ollama error (status {response.status_code})"
                try:
                    error_detail = response.json().get('error', '')
                    if error_detail:
                        error_msg += f": {error_detail}"
                except:
                    pass
                return f"Error: {error_msg}"
                
        except requests.exceptions.ConnectionError:
            return "Error: Cannot connect to Ollama. Please run 'ollama serve'."
        except requests.exceptions.Timeout:
            return "Error: Request timed out. Please try again."
        except Exception as e:
            return f"Error: {str(e)}"

# Initialize RAG pipeline
rag_pipeline = RAGPipeline()

# API Endpoints
@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'message': 'StarRAG Bot API is running',
        'version': '1.0'
    })

@app.route('/api/upload', methods=['POST'])
def upload_file():
    """Handle file upload"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if file and rag_pipeline.allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        
        success = rag_pipeline.process_document(file_path, filename)
        
        os.remove(file_path)
        
        if success:
            return jsonify({'message': f'File {filename} processed successfully'})
        else:
            return jsonify({'error': 'Failed to process document'}), 500
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/api/documents', methods=['GET'])
def get_documents():
    """Get list of uploaded documents"""
    if documents_collection is None:
        return jsonify({
            'documents': [],
            'message': 'MongoDB not available'
        })
    
    try:
        docs = list(documents_collection.find({}, {'filename': 1, 'created_at': 1}))
        return jsonify({
            'documents': [
                {
                    'id': str(doc['_id']),
                    'filename': doc['filename'],
                    'created_at': doc['created_at'].isoformat()
                }
                for doc in docs
            ]
        })
    except Exception as e:
        return jsonify({
            'documents': [],
            'error': str(e)
        })

@app.route('/api/documents/<doc_id>', methods=['DELETE'])
def delete_document(doc_id):
    """Delete a document"""
    if documents_collection is None:
        return jsonify({'error': 'MongoDB not available'}), 500
    
    try:
        # Remove document from MongoDB
        result = documents_collection.delete_one({'_id': doc_id})
        
        if result.deleted_count == 0:
            return jsonify({'error': 'Document not found'}), 404
            
        # Rebuild FAISS index without this document
        global faiss_index, document_chunks
        
        if faiss_index is not None:
            all_docs = list(documents_collection.find())
            
            faiss_index = None
            document_chunks = []
            
            if all_docs:
                all_embeddings = []
                
                for doc in all_docs:
                    for chunk in doc.get("chunks", []):
                        embedding = np.array(chunk["embedding"])
                        all_embeddings.append(embedding)
                        document_chunks.append({
                            "text": chunk["text"],
                            "metadata": chunk["metadata"],
                            "doc_id": doc["_id"]
                        })
                
                if all_embeddings:
                    embeddings_array = np.vstack(all_embeddings)
                    dimension = embeddings_array.shape[1]
                    faiss_index = faiss.IndexFlatL2(dimension)
                    faiss_index.add(embeddings_array.astype('float32'))
        
        return jsonify({'message': 'Document deleted successfully'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/models', methods=['GET'])
def get_models():
    """Get available Ollama models"""
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=15)
        if response.status_code == 200:
            models = response.json().get('models', [])
            model_names = [model['name'] for model in models]
            
            # Filter to only show allowed models that are available
            available_models = [m for m in ALLOWED_MODELS if m in model_names]
            
            if available_models:
                return jsonify({
                    'models': available_models,
                    'status': 'connected'
                })
            else:
                return jsonify({
                    'models': ALLOWED_MODELS,
                    'status': 'no_models',
                    'message': f'None of the allowed models found. Available models: {model_names}'
                })
        else:
            return jsonify({
                'models': ALLOWED_MODELS,
                'status': 'error',
                'message': f'Ollama responded with status {response.status_code}'
            })
    except requests.exceptions.ConnectionError:
        return jsonify({
            'models': ALLOWED_MODELS,
            'status': 'disconnected',
            'message': 'Cannot connect to Ollama. Run "ollama serve" in terminal.'
        })
    except Exception as e:
        return jsonify({
            'models': ALLOWED_MODELS,
            'status': 'error',
            'message': str(e)
        })

@app.route('/api/query', methods=['POST'])
def query():
    """Handle user queries"""
    data = request.json
    user_query = data.get('query', '')
    model = data.get('model', 'gemma3:1b')
    conversation_id = data.get('conversation_id')
    
    if not user_query:
        return jsonify({'error': 'No query provided'}), 400
    
    # Validate model
    if model not in ALLOWED_MODELS:
        return jsonify({'error': f'Model {model} is not allowed'}), 400
    
    # Search for relevant chunks
    relevant_chunks = rag_pipeline.similarity_search(user_query, k=5)
    
    if not relevant_chunks:
        return jsonify({
            'response': 'I have no relevant information. Please upload documents first.',
            'sources': [],
            'conversation_id': conversation_id
        })
    
    # Generate response
    response = rag_pipeline.generate_response(user_query, relevant_chunks, model)
    
    # Create or update conversation
    if not conversation_id:
        conversation_id = str(uuid.uuid4())
    
    conversation_data = {
        'conversation_id': conversation_id,
        'query': user_query,
        'response': response,
        'model': model,
        'sources': list(set([chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks])),
        'timestamp': datetime.utcnow()
    }
    
    if conversations_collection is not None:
        try:
            conversations_collection.insert_one(conversation_data)
        except Exception as e:
            print(f"Warning: Could not store conversation: {e}")
    
    return jsonify({
        'response': response,
        'sources': list(set([chunk['metadata'].get('source', 'Unknown') for chunk in relevant_chunks])),
        'conversation_id': conversation_id
    })

@app.route('/api/conversations', methods=['GET'])
def list_conversations():
    """List all conversations"""
    if conversations_collection is None:
        return jsonify({'conversations': [], 'message': 'MongoDB not available'})
    
    try:
        pipeline = [
            {"$sort": {"timestamp": 1}},
            {"$group": {
                "_id": "$conversation_id",
                "first_query": {"$first": "$query"},
                "first_timestamp": {"$first": "$timestamp"},
                "model": {"$first": "$model"}
            }},
            {"$sort": {"first_timestamp": -1}}
        ]
        conversations = list(conversations_collection.aggregate(pipeline))
        return jsonify({
            'conversations': [
                {
                    'conversation_id': c['_id'],
                    'first_query': c.get('first_query', ''),
                    'first_timestamp': c.get('first_timestamp').isoformat() if c.get('first_timestamp') else '',
                    'model': c.get('model', '')
                }
                for c in conversations
            ]
        })
    except Exception as e:
        return jsonify({'conversations': [], 'error': str(e)})

@app.route('/api/conversations/<conversation_id>', methods=['GET'])
def get_conversation(conversation_id):
    """Get conversation history"""
    if conversations_collection is None:
        return jsonify({'conversations': [], 'message': 'MongoDB not available'})
    
    try:
        conversations = list(conversations_collection.find(
            {'conversation_id': conversation_id},
            {'_id': 0}
        ).sort('timestamp', 1))
        
        return jsonify({'conversations': conversations})
    except Exception as e:
        return jsonify({'conversations': [], 'error': str(e)})

if __name__ == '__main__':
    # Load existing documents on startup
    if documents_collection is not None:
        try:
            docs = list(documents_collection.find())
            if docs:
                all_embeddings = []
                document_chunks = []
                
                for doc in docs:
                    for chunk in doc.get("chunks", []):
                        embedding = np.array(chunk["embedding"])
                        all_embeddings.append(embedding)
                        document_chunks.append({
                            "text": chunk["text"],
                            "metadata": chunk["metadata"],
                            "doc_id": doc["_id"]
                        })
                
                if all_embeddings:
                    embeddings_array = np.vstack(all_embeddings)
                    dimension = embeddings_array.shape[1]
                    faiss_index = faiss.IndexFlatL2(dimension)
                    faiss_index.add(embeddings_array.astype('float32'))
                    print(f"✅ Loaded {len(docs)} documents with {len(document_chunks)} chunks")
        except Exception as e:
            print(f"⚠️ Error loading existing documents: {e}")
    
    print("🚀 StarRAG Bot API starting...")
    app.run(debug=True, port=5000)